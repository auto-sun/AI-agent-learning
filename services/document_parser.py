from pathlib import Path

from docx import Document
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


def extract_text_from_txt(file_path: str) -> str:
    """
    提取 txt 文件文本。
    """
    path = Path(file_path)

    return path.read_text(encoding="utf-8")


def extract_text_from_md(file_path: str) -> str:
    """
    提取 markdown 文件文本。

    md 本质上也是文本文件，所以可以先按普通文本读取。
    后面如果你想去掉 #、*、``` 等 markdown 标记，可以再做清洗。
    """
    path = Path(file_path)

    return path.read_text(encoding="utf-8")


def extract_text_from_pdf(file_path: str) -> str:
    """
    提取 PDF 文件文本。

    注意：
    1. 适合普通文字型 PDF
    2. 不适合扫描版 PDF
    3. 扫描版 PDF 需要 OCR，后面再学
    """
    reader = PdfReader(file_path)

    page_texts = []

    for page_index, page in enumerate(reader.pages):
        text = page.extract_text()

        if text and text.strip():
            page_texts.append(
                f"\n\n--- 第 {page_index + 1} 页 ---\n\n{text.strip()}"
            )

    return "\n".join(page_texts).strip()


def extract_text_from_docx(file_path: str) -> str:
    """
    提取 docx 文件文本。

    先提取段落文本。
    表格内容这节课先做简单提取。
    """
    document = Document(file_path)

    texts = []

    # 1. 提取普通段落
    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text.strip()

        if paragraph_text:
            texts.append(paragraph_text)

    # 2. 简单提取表格文本
    for table_index, table in enumerate(document.tables):
        texts.append(f"\n--- 表格 {table_index + 1} ---")

        for row in table.rows:
            row_texts = []

            for cell in row.cells:
                cell_text = cell.text.strip()

                if cell_text:
                    row_texts.append(cell_text)

            if row_texts:
                texts.append(" | ".join(row_texts))

    return "\n".join(texts).strip()


def extract_text_from_file(file_path: str) -> dict:
    """
    根据文件后缀自动选择解析方式。

    返回 dict，而不是只返回 text，是为了把解析信息也返回给前端。
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"暂不支持 {suffix} 文件。当前只支持：txt、md、pdf、docx"
        )

    if suffix == ".txt":
        text = extract_text_from_txt(file_path)
        file_type = "txt"
    elif suffix == ".md":
        text = extract_text_from_md(file_path)
        file_type = "md"
    elif suffix == ".pdf":
        text = extract_text_from_pdf(file_path)
        file_type = "pdf"
    elif suffix == ".docx":
        text = extract_text_from_docx(file_path)
        file_type = "docx"
    else:
        raise ValueError(f"暂不支持的文件类型：{suffix}")

    clean_text = text.strip()

    if not clean_text:
        raise ValueError(
            "文件解析后没有提取到有效文本。"
            "如果是扫描版 PDF，需要后续接 OCR。"
        )

    return {
        "text": clean_text,
        "file_type": file_type,
        "file_suffix": suffix,
        "text_length": len(clean_text),
    }